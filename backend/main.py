# backend/main.py
import os
import uuid
from datetime import datetime, timedelta
from fastapi import FastAPI, UploadFile, File, Form, HTTPException, Query
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, FileResponse
from pydantic import BaseModel
import aiofiles
from typing import Optional, Dict, Any, List
from collections import defaultdict
import re
import json
import math

# New imports for OpenAI endpoint (if you use)
import openai

# Additional imports for new quick-action features / reports
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import io
import threading

app = FastAPI(title="AI Compliance Assistant (demo)")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

UPLOAD_DIR = os.path.join(os.path.dirname(__file__), "uploads")
os.makedirs(UPLOAD_DIR, exist_ok=True)

METRICS_TREND: List[Dict[str, Any]] = []
DISTRIBUTION_COUNTER: Dict[str, int] = {}
CATEGORY_COUNTER: Dict[str, int] = {}

def init_synthetic_metrics():
    now = datetime.utcnow()
    for i in range(7, 0, -1):
        t = now - timedelta(days=i)
        compliance = round(78 + (i % 5) * 1.8 + (i % 3) * 0.4, 1)
        risk_items = max(2, int(120 - compliance * 1.1 + (i % 4) * 6))
        remediated = max(0, int(risk_items * (0.45 + ((i % 3) * 0.08))))
        METRICS_TREND.append({
            "timestamp": t.isoformat() + "Z",
            "complianceScore": compliance,
            "riskItems": risk_items,
            "remediatedItems": remediated
        })
    global DISTRIBUTION_COUNTER, CATEGORY_COUNTER
    DISTRIBUTION_COUNTER = {"Slack": 35, "Email": 28, "GitHub": 18, "Jira": 12, "Database": 7}
    CATEGORY_COUNTER = {"SSN": 45, "Credit Card": 32, "Email": 89, "Phone": 67, "Address": 54, "Name": 123}

init_synthetic_metrics()

class ChatRequest(BaseModel):
    message: Optional[str] = ""

# --- PII helpers (kept from your file) ---
def pii_detector_agent(text: str) -> Dict[str, Any]:
    # minimal demo detector - you likely replaced with improved detector earlier
    emails = re.findall(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}", text)
    phones = re.findall(r"\b(?:\+?\d{1,3}[-.\s]?)?(?:\d{2,4}[-.\s]?){2,4}\d{2,4}\b", text)
    ssn_like = re.findall(r"\b\d{3}-\d{2}-\d{4}\b", text)
    return {"emails": emails, "phones": phones, "ssn_like": ssn_like}

def anonymizer_agent(text: str) -> str:
    out = re.sub(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}", "[REDACTED_EMAIL]", text)
    out = re.sub(r"\b\d{3}-\d{2}-\d{4}\b", "[REDACTED_SSN]", out)
    return out[:2000]

def summarizer_agent(text: str) -> str:
    return f"Demo summary: {text[:240]}"

def orchestrate_text_message(message: str) -> Dict[str, Any]:
    if not message:
        return {"reply": "Hi — how can I help? You can upload a file or ask me to scan for PII."}
    low = message.lower()
    if "anonymize" in low or "mask" in low or "redact" in low:
        return {"reply": "Anonymized (demo): " + anonymizer_agent(message)}
    if "scan" in low or "pii" in low or "sensitive" in low:
        det = pii_detector_agent(message)
        return {
            "reply": f"PII check (demo): emails={len(det['emails'])}, phones={len(det['phones'])}, ssn_like={len(det['ssn_like'])}.",
            "pii": det
        }
    return {"reply": summarizer_agent(message)}

@app.post("/chat")
async def chat_endpoint(req: ChatRequest):
    out = orchestrate_text_message(req.message or "")
    return JSONResponse(content=out)

# OpenAI / ai_predict kept as in your file
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")

class PredictRequest(BaseModel):
    prompt: str
    max_tokens: int = 512
    temperature: float = 0.6

@app.post("/ai/predict")
async def ai_predict(req: PredictRequest):
    if not OPENAI_API_KEY:
        return JSONResponse(status_code=500, content={"error": "OpenAI key not set on server."})
    try:
        openai.api_key = OPENAI_API_KEY
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "You are a professional compliance assistant."},
                {"role": "user", "content": req.prompt}
            ],
            max_tokens=req.max_tokens,
            temperature=req.temperature
        )
        text = ""
        if getattr(response, "choices", None):
            try:
                text = response.choices[0].message.get("content", "")
            except Exception:
                text = response.choices[0].get("text", "") if isinstance(response.choices[0], dict) else ""
        return JSONResponse(content={"result": text})
    except Exception as e:
        return JSONResponse(status_code=500, content={"error": str(e)})

# --- compute_metric_from_findings unchanged ---
def compute_metric_from_findings(findings: Dict[str, List[str]]) -> Dict[str, Any]:
    # this function expects structured findings; keep as-is for demo
    emails = 0
    phones = 0
    ssn_like = 0
    if isinstance(findings, dict):
        # try to access commonly named keys if present
        if "emails" in findings:
            emails = len(findings.get("emails") or [])
        elif "counts" in findings and isinstance(findings["counts"], dict):
            emails = int(findings["counts"].get("EMAIL", 0) or findings["counts"].get("emails", 0) or 0)
        if "phones" in findings:
            phones = len(findings.get("phones") or [])
        elif "counts" in findings and isinstance(findings["counts"], dict):
            phones = int(findings["counts"].get("PHONE", 0) or findings["counts"].get("phones", 0) or 0)
        if "ssn_like" in findings:
            ssn_like = len(findings.get("ssn_like") or [])
        elif "counts" in findings and isinstance(findings["counts"], dict):
            ssn_like = int(findings["counts"].get("SSN", 0) or findings["counts"].get("ssn_like", 0) or 0)
    risk_items = emails + phones + ssn_like
    base = 90
    penalty = risk_items * 3
    compliance = max(25, base - penalty)
    remediated = int(risk_items * 0.5)
    return {"complianceScore": round(float(compliance), 1), "riskItems": int(risk_items), "remediatedItems": int(remediated)}

# -------------------- Upload endpoint (kept close to original) --------------------
@app.post("/chat/upload")
async def chat_upload(file: UploadFile = File(...), message: str = Form(None)):
    uid = uuid.uuid4().hex
    safe_name = f"{uid}_{file.filename}"
    dest = os.path.join(UPLOAD_DIR, safe_name)
    try:
        async with aiofiles.open(dest, "wb") as f:
            while True:
                chunk = await file.read(1024 * 1024)
                if not chunk:
                    break
                await f.write(chunk)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

    # try to read as text
    try:
        async with aiofiles.open(dest, "r", encoding="utf-8", errors="ignore") as fr:
            content = await fr.read()
    except Exception:
        content = ""
    if not content:
        content = f"(binary or non-text file) {file.filename}"

    pii = pii_detector_agent(content)
    summary = summarizer_agent(content)
    reply = f"Processed {file.filename} — summary: {summary}"

    global DISTRIBUTION_COUNTER, CATEGORY_COUNTER, METRICS_TREND
    if pii.get("emails"):
        DISTRIBUTION_COUNTER["Email"] = DISTRIBUTION_COUNTER.get("Email", 0) + len(pii["emails"])
        CATEGORY_COUNTER["Email"] = CATEGORY_COUNTER.get("Email", 0) + len(pii["emails"])
    if pii.get("phones"):
        DISTRIBUTION_COUNTER["Slack"] = DISTRIBUTION_COUNTER.get("Slack", 0) + len(pii["phones"])
        CATEGORY_COUNTER["Phone"] = CATEGORY_COUNTER.get("Phone", 0) + len(pii["phones"])
    if pii.get("ssn_like"):
        CATEGORY_COUNTER["SSN"] = CATEGORY_COUNTER.get("SSN", 0) + len(pii["ssn_like"])
        DISTRIBUTION_COUNTER["Database"] = DISTRIBUTION_COUNTER.get("Database", 0) + len(pii["ssn_like"])

    metric = compute_metric_from_findings(pii)
    metric_point = {
        "timestamp": datetime.utcnow().isoformat() + "Z",
        "complianceScore": metric["complianceScore"],
        "riskItems": metric["riskItems"],
        "remediatedItems": metric["remediatedItems"],
    }
    METRICS_TREND.append(metric_point)

    return JSONResponse(content={"reply": reply, "report": {"summary": summary, "pii": pii}})

# -------------------- Aggregated metrics endpoint (hourly) --------------------
def parse_iso_z(ts: str) -> Optional[datetime]:
    if not ts:
        return None
    try:
        # strip trailing Z and parse
        if ts.endswith("Z"):
            ts = ts[:-1]
        return datetime.fromisoformat(ts)
    except Exception:
        # fallback try other formats
        try:
            return datetime.strptime(ts, "%Y-%m-%dT%H:%M:%S")
        except Exception:
            return None

@app.get("/metrics")
async def get_metrics(range_hours: int = Query(24, ge=1, le=24*30)):
    """
    Return hourly-aggregated metrics for the last `range_hours`.
    Aggregation: average complianceScore, sum riskItems/remediatedItems per hour bucket.
    """
    now = datetime.utcnow()
    start_ts = now - timedelta(hours=range_hours)

    # bucket by hour start
    buckets = defaultdict(list)
    for p in METRICS_TREND:
        ts = p.get("timestamp")
        dt = parse_iso_z(ts)
        if not dt:
            continue
        if dt < start_ts:
            continue
        hour_start = dt.replace(minute=0, second=0, microsecond=0)
        buckets[hour_start.isoformat() + "Z"].append(p)

    trend = []
    for hour_key in sorted(buckets.keys()):
        vals = buckets[hour_key]
        if not vals:
            continue
        avg_compliance = sum(v.get("complianceScore", 0) for v in vals) / len(vals)
        sum_risk = sum(v.get("riskItems", 0) for v in vals)
        sum_rem = sum(v.get("remediatedItems", 0) for v in vals)
        trend.append({
            "timestamp": hour_key,
            "complianceScore": round(avg_compliance, 1),
            "riskItems": int(sum_risk),
            "remediatedItems": int(sum_rem)
        })

    # If no data points (fresh server), return last synthetic points as fallback
    if not trend and METRICS_TREND:
        # simply return last N points downsampled into hours roughly
        fallback = METRICS_TREND[-min(len(METRICS_TREND), range_hours):]
        trend = fallback

    # distribution/categories unchanged from earlier logic
    distribution = []
    for k, v in DISTRIBUTION_COUNTER.items():
        color = "#0EA5E9" if "Slack" in k else ("#059669" if "Email" in k else "#D97706")
        distribution.append({"name": k, "value": v, "color": color})
    categories = []
    for k, v in CATEGORY_COUNTER.items():
        severity = "Critical" if k.lower() in ("ssn", "credit card") else ("Medium" if v > 40 else "Low")
        categories.append({"category": k, "count": v, "severity": severity})

    return JSONResponse(content={
        "trend": trend,
        "distribution": distribution,
        "categories": categories,
        "last_updated": now.isoformat() + "Z"
    })

# -------------------- Export CSV uses same aggregation (range_hours param supported) --------------------
@app.get("/metrics/export/csv")
async def export_metrics_csv(range_hours: int = Query(24, ge=1, le=24*30)):
    # Use the same aggregation logic as /metrics
    now = datetime.utcnow()
    start_ts = now - timedelta(hours=range_hours)
    buckets = defaultdict(list)
    for p in METRICS_TREND:
        ts = p.get("timestamp")
        dt = parse_iso_z(ts)
        if not dt:
            continue
        if dt < start_ts:
            continue
        hour_start = dt.replace(minute=0, second=0, microsecond=0)
        buckets[hour_start.isoformat() + "Z"].append(p)

    trend = []
    for hour_key in sorted(buckets.keys()):
        vals = buckets[hour_key]
        if not vals:
            continue
        avg_compliance = sum(v.get("complianceScore", 0) for v in vals) / len(vals)
        sum_risk = sum(v.get("riskItems", 0) for v in vals)
        sum_rem = sum(v.get("remediatedItems", 0) for v in vals)
        trend.append({
            "timestamp": hour_key,
            "complianceScore": round(avg_compliance, 1),
            "riskItems": int(sum_risk),
            "remediatedItems": int(sum_rem)
        })

    # fallback
    if not trend and METRICS_TREND:
        trend = METRICS_TREND[-min(len(METRICS_TREND), range_hours):]

    csv_path = os.path.join(UPLOAD_DIR, "metrics_export.csv")
    try:
        with open(csv_path, "w", encoding="utf-8") as fh:
            fh.write("timestamp,complianceScore,riskItems,remediatedItems\n")
            for p in trend:
                fh.write(f'{p["timestamp"]},{p["complianceScore"]},{p["riskItems"]},{p["remediatedItems"]}\n')
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    return FileResponse(csv_path, media_type="text/csv", filename="metrics_export.csv")

# -------------------- Quick Actions (NEW endpoints appended) --------------------

# last scan storage
LAST_SCAN: Dict[str, Any] = {"status": "idle", "timestamp": None, "findings": None, "summary": None, "remediation": []}

@app.post("/scan/start")
async def start_scan(data_source: Optional[str] = Form(None)):
    """
    Start a simulated scan in the background.
    Accepts optional data_source (string).
    """
    global LAST_SCAN, METRICS_TREND, DISTRIBUTION_COUNTER, CATEGORY_COUNTER

    # mark running
    LAST_SCAN = {"status": "running", "timestamp": datetime.utcnow().isoformat() + "Z", "findings": None, "summary": None, "remediation": []}

    def scan_job(ds):
        # simulate variable results depending on ds
        found = {}
        if ds and "email" in ds.lower():
            found["emails"] = ["alice@example.com", "bob@example.org"]
            found["phones"] = []
            found["ssn_like"] = []
        elif ds and "slack" in ds.lower():
            found["emails"] = []
            found["phones"] = ["+1-555-123-4567"]
            found["ssn_like"] = ["123-45-6789"]
        else:
            # mixed sample
            found["emails"] = ["user1@company.com"]
            found["phones"] = ["+44 7700 900123"]
            found["ssn_like"] = []

        # update counters
        if found.get("emails"):
            DISTRIBUTION_COUNTER["Email"] = DISTRIBUTION_COUNTER.get("Email", 0) + len(found["emails"])
            CATEGORY_COUNTER["Email"] = CATEGORY_COUNTER.get("Email", 0) + len(found["emails"])
        if found.get("phones"):
            DISTRIBUTION_COUNTER["Slack"] = DISTRIBUTION_COUNTER.get("Slack", 0) + len(found["phones"])
            CATEGORY_COUNTER["Phone"] = CATEGORY_COUNTER.get("Phone", 0) + len(found["phones"])
        if found.get("ssn_like"):
            CATEGORY_COUNTER["SSN"] = CATEGORY_COUNTER.get("SSN", 0) + len(found["ssn_like"])
            DISTRIBUTION_COUNTER["Database"] = DISTRIBUTION_COUNTER.get("Database", 0) + len(found["ssn_like"])

        metric = compute_metric_from_findings(found)
        metric_point = {
            "timestamp": datetime.utcnow().isoformat() + "Z",
            "complianceScore": metric["complianceScore"],
            "riskItems": metric["riskItems"],
            "remediatedItems": metric["remediatedItems"],
        }
        METRICS_TREND.append(metric_point)

        remediation_plans = []
        if found.get("emails"):
            remediation_plans.append("Remove public emails or mask them")
        if found.get("phones"):
            remediation_plans.append("Encrypt phone numbers and restrict access")
        if found.get("ssn_like"):
            remediation_plans.append("Move SSNs to secure vault and anonymize")

        LAST_SCAN.update({
            "status": "completed",
            "timestamp": datetime.utcnow().isoformat() + "Z",
            "findings": found,
            "summary": f"Found {sum(len(v) for v in found.values())} PII items",
            "remediation": remediation_plans
        })

    t = threading.Thread(target=scan_job, args=(data_source,), daemon=True)
    t.start()

    return JSONResponse(content={"status": "started", "message": "Scan started in background"})

@app.get("/scan/status")
async def scan_status():
    """
    Return the last scan status and findings (useful for Quick Actions 'View Risk Details').
    """
    return JSONResponse(content=LAST_SCAN)

@app.post("/remediation/apply")
async def remediation_apply(plan_index: Optional[int] = Form(None)):
    """
    Simulate applying a remediation plan returned by last scan.
    plan_index: index of remediation in LAST_SCAN.remediation array
    """
    global LAST_SCAN
    if not LAST_SCAN or not LAST_SCAN.get("remediation"):
        return JSONResponse(content={"status": "no_plans", "message": "No remediation plans available"})
    if plan_index is None:
        # apply all (simulated)
        applied = LAST_SCAN["remediation"]
    else:
        try:
            applied = [LAST_SCAN["remediation"][int(plan_index)]]
        except Exception:
            return JSONResponse(content={"status": "error", "message": "Invalid plan_index"})
    # update metrics to reflect remediation effect (simple simulation: reduce riskItems)
    # We'll append a metric point showing some remediated items
    rem_count = len(applied)
    new_point = {
        "timestamp": datetime.utcnow().isoformat() + "Z",
        "complianceScore": min(100, (METRICS_TREND[-1]["complianceScore"] + rem_count * 0.5) if METRICS_TREND else 90),
        "riskItems": max(0, (METRICS_TREND[-1]["riskItems"] - rem_count) if METRICS_TREND else 0),
        "remediatedItems": (METRICS_TREND[-1]["remediatedItems"] + rem_count) if METRICS_TREND else rem_count
    }
    METRICS_TREND.append(new_point)
    return JSONResponse(content={"status": "applied", "applied": applied, "new_metric_point": new_point})

# Report generation helpers (simple structured docx/pdf)
def build_docx_report_structured(payload: dict, title: str = "Compliance Report") -> str:
    doc = Document()
    doc.add_heading(title, level=1)

    # 1. PII detected
    doc.add_heading("1. PII detected", level=2)
    pii = payload.get("pii", {})
    if pii:
        for k, v in pii.items():
            doc.add_paragraph(f"{k}: {len(v)} occurrences")
            # list up to 10 items for readability
            if isinstance(v, list):
                for item in v[:10]:
                    doc.add_paragraph(f" - {item}")
    else:
        doc.add_paragraph("No PII detected.")

    # 2. Scanning done
    doc.add_heading("2. Scanning done", level=2)
    doc.add_paragraph(payload.get("summary", "Scan completed."))

    # 3. Risks extracted (number)
    risks = payload.get("risks", [])
    doc.add_heading(f"3. Risks extracted (number): {len(risks)}", level=2)

    # 4. List of risks
    doc.add_heading("4. List of risks", level=2)
    if risks:
        for r in risks:
            doc.add_paragraph(f"- {r}")
    else:
        doc.add_paragraph("- No explicit risks flagged.")

    # 5. Remediation plans generated(number)
    remediation = payload.get("remediation", [])
    doc.add_heading(f"5. Remediation plans generated: {len(remediation)}", level=2)

    # 6. List of plans
    doc.add_heading("6. List of plans", level=2)
    if remediation:
        for r in remediation:
            doc.add_paragraph(f"- {r}")
    else:
        doc.add_paragraph("- No remediation plans generated.")

    out_path = os.path.join(UPLOAD_DIR, f"report_{uuid.uuid4().hex}.docx")
    doc.save(out_path)
    return out_path

def build_pdf_report_structured(payload: dict, title: str = "Compliance Report") -> str:
    out_path = os.path.join(UPLOAD_DIR, f"report_{uuid.uuid4().hex}.pdf")
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    y = height - 40
    c.setFont("Helvetica-Bold", 16)
    c.drawString(40, y, title)
    y -= 30
    c.setFont("Helvetica", 11)

    c.drawString(40, y, "1. PII detected")
    y -= 18
    pii = payload.get("pii", {})
    if pii:
        for k, v in pii.items():
            c.drawString(50, y, f"{k}: {len(v)}")
            y -= 14
            if isinstance(v, list):
                for item in v[:5]:
                    c.drawString(60, y, f"- {item}")
                    y -= 12
                    if y < 70:
                        c.showPage()
                        y = height - 40
    else:
        c.drawString(50, y, "No PII detected.")
        y -= 18

    c.drawString(40, y, "2. Scanning done")
    y -= 14
    c.drawString(50, y, payload.get("summary", "Scan completed."))
    y -= 22

    c.drawString(40, y, f"3. Risks extracted (number): {len(payload.get('risks', []))}")
    y -= 18
    c.drawString(40, y, "4. List of risks")
    y -= 14
    for r in payload.get("risks", []):
        c.drawString(50, y, f"- {r}")
        y -= 12
        if y < 70:
            c.showPage()
            y = height - 40

    c.drawString(40, y, f"5. Remediation plans generated: {len(payload.get('remediation', []))}")
    y -= 18
    c.drawString(40, y, "6. List of plans")
    y -= 14
    for r in payload.get("remediation", []):
        c.drawString(50, y, f"- {r}")
        y -= 12
        if y < 70:
            c.showPage()
            y = height - 40

    c.save()
    buffer.seek(0)
    with open(out_path, "wb") as fh:
        fh.write(buffer.read())
    return out_path

@app.post("/reports/generate")
async def generate_report(format: Optional[str] = Form("docx"), title: Optional[str] = Form("Compliance Report"), payload_json: Optional[str] = Form("{}")):
    """
    Generate a structured report (docx or pdf).
    Accepts form fields:
      - format: "docx" or "pdf"
      - title: string
      - payload_json: json string with keys: pii (dict), risks (list), remediation (list), summary (str)
    """
    try:
        payload = json.loads(payload_json or "{}")
    except Exception:
        payload = {}

    if format not in ("docx", "pdf"):
        raise HTTPException(status_code=400, detail="format must be 'docx' or 'pdf'")

    if format == "docx":
        out = build_docx_report_structured(payload, title=title)
    else:
        out = build_pdf_report_structured(payload, title=title)

    return FileResponse(out, media_type="application/octet-stream", filename=os.path.basename(out))

# -------------------- Compatibility wrappers & activity feed (append below existing file) --------------------
from fastapi import Request, Body
from typing import Any as TAny

# Simple in-memory activity log (kept minimal for demo)
ACTIVITY: List[Dict[str, TAny]] = []

def _record_activity(kind: str, title: str, details: dict):
    """Append a small activity record (kept very simple)."""
    try:
        ACTIVITY.insert(0, {
            "id": uuid.uuid4().hex,
            "kind": kind,
            "title": title,
            "details": details,
            "timestamp": datetime.utcnow().isoformat() + "Z"
        })
        # keep the list reasonably short
        if len(ACTIVITY) > 200:
            ACTIVITY.pop()
    except Exception:
        pass

@app.post("/scan/start_json")
async def start_scan_json(request: Request):
    """
    Compatibility wrapper: accepts application/json { "data_source": "..." } or { "source": "..." }.
    Internally re-uses your existing start_scan(data_source=...) function.
    """
    try:
        payload = await request.json()
    except Exception:
        payload = {}
    ds = payload.get("data_source") or payload.get("source") or payload.get("dataSource") or None
    # call existing start_scan (it returns a JSONResponse)
    resp = await start_scan(data_source=ds)
    # record a small activity
    _record_activity("scan_started", "Manual scan started", {"data_source": ds})
    return resp

@app.post("/remediation/apply_json")
async def remediation_apply_json(request: Request):
    """
    Accepts JSON: { "plan_index": 0 } or { "apply_all": true }.
    Calls the existing remediation_apply(plan_index=...) function.
    """
    try:
        payload = await request.json()
    except Exception:
        payload = {}
    plan_index = payload.get("plan_index", None)
    # If client asks to apply all via flag, leave plan_index None
    if payload.get("apply_all") is True:
        plan_index = None
    resp = await remediation_apply(plan_index=plan_index)
    _record_activity("remediation_applied", "Remediation applied (wrapper)", {"plan_index": plan_index})
    return resp

@app.post("/reports/generate_json")
async def generate_report_json(format: Optional[str] = Query("docx"), payload: Optional[dict] = Body(None), title: Optional[str] = Query("Compliance Report")):
    """
    Compatibility wrapper: POST JSON body to create a report.
    Query param 'format' controls docx/pdf. JSON body keys: pii, risks, remediation, summary.
    Example POST body:
      { "pii": {"emails": ["a@b.com"]}, "risks": ["something"], "remediation": ["plan1"], "summary": "..." }
    Returns FileResponse (same as /reports/generate).
    """
    payload = payload or {}
    # delegate to same builder functions used by original endpoint
    if format not in ("docx", "pdf"):
        raise HTTPException(status_code=400, detail="format must be 'docx' or 'pdf'")

    if format == "docx":
        out = build_docx_report_structured(payload, title=title)
    else:
        out = build_pdf_report_structured(payload, title=title)

    # record activity
    _record_activity("report_generated", "Report generated (JSON wrapper)", {"format": format, "file": os.path.basename(out)})
    return FileResponse(out, media_type="application/octet-stream", filename=os.path.basename(out))

@app.get("/activity")
async def get_activity(limit: int = Query(25, ge=1, le=200)):
    """
    Simple activity feed for the UI. Returns the most recent activity records (in-memory).
    This is ephemeral (server restart clears it) — useful for demo / UI integration.
    """
    # format quick human-friendly time label
    def _fmt(item):
        ts = item.get("timestamp")
        try:
            dt = parse_iso_z(ts)
            if dt:
                diff = datetime.utcnow() - dt
                mins = int(diff.total_seconds() // 60)
                if mins < 60:
                    label = f"{mins}m ago"
                else:
                    hrs = mins // 60
                    label = f"{hrs}h ago"
            else:
                label = ts
        except Exception:
            label = ts
        return {
            "id": item.get("id"),
            "kind": item.get("kind"),
            "title": item.get("title"),
            "details": item.get("details"),
            "timestamp": item.get("timestamp"),
            "timeLabel": label
        }

    return JSONResponse(content=[_fmt(it) for it in ACTIVITY[:limit]])

# For convenience: make the original endpoints also record some activity (without modifying their logic).
# We add lightweight wrappers that call the original functions and record events.
# NOTE: we don't change the original functions — these wrappers are optional helpers for the UI.

@app.post("/scan/start_and_record")
async def start_scan_and_record(data_source: Optional[str] = Form(None)):
    """
    Backwards-compatible wrapper that behaves like /scan/start but also records an activity.
    Use this from the UI if you want an activity entry guaranteed.
    """
    resp = await start_scan(data_source=data_source)
    try:
        _record_activity("scan_started", "Scan started (form wrapper)", {"data_source": data_source})
    except Exception:
        pass
    return resp

@app.post("/reports/generate_and_record")
async def generate_report_and_record(format: Optional[str] = Form("docx"), title: Optional[str] = Form("Compliance Report"), payload_json: Optional[str] = Form("{}")):
    """
    Backwards-compatible wrapper for form-based POSTs that also records the event.
    Delegates to existing generate_report behavior and records an activity entry.
    """
    resp = await generate_report(format=format, title=title, payload_json=payload_json)
    try:
        _record_activity("report_generated", "Report generated (form wrapper)", {"format": format})
    except Exception:
        pass
    return resp

# ---------------------------------------------------------------------------------------
# End of appended compatibility helpers
# ---------------------------------------------------------------------------------------
